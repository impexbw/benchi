"""Universal file reader tool.

Reads uploaded files from Frappe, parses them by type, and returns structured
content the AI can reason about. Supports TXT, CSV, Excel (XLSX/XLS), PDF,
and JSON files.
"""
import frappe
from erpnext_ai_bots.tools.base import BaseTool


class ReadFileTool(BaseTool):
    name = "core.read_file"
    description = (
        "Read and parse an uploaded file. Supports TXT, CSV, Excel (XLSX/XLS), Word (DOCX), "
        "PDF, and JSON files. Returns the file content as structured data that "
        "you can analyze, cross-reference with ERPNext, or act on. "
        "Use this when a user uploads a non-image file and asks questions about it. "
        "For images, use core_analyze_image instead."
    )
    parameters = {
        "file_url": {
            "type": "string",
            "description": "The file URL from the upload (e.g. /private/files/data.csv)",
        },
        "max_rows": {
            "type": "integer",
            "description": "Max rows to return for CSV/Excel files (default 100, max 500)",
        },
    }
    required_params = ["file_url"]
    action_type = "Read"
    required_ptype = None

    def execute(self, file_url, max_rows=100, **kwargs):
        max_rows = min(int(max_rows), 500)

        content_bytes = self._get_file_content(file_url)
        if content_bytes is None:
            return {"error": f"Could not read file at {file_url}"}

        ext = file_url.rsplit(".", 1)[-1].lower() if "." in file_url else ""

        if ext in ("txt", "log", "md", "text"):
            return self._read_text(content_bytes)
        elif ext == "json":
            return self._read_json(content_bytes)
        elif ext in ("csv", "tsv"):
            return self._read_csv(content_bytes, max_rows)
        elif ext in ("xlsx", "xls"):
            return self._read_excel(content_bytes, max_rows)
        elif ext == "pdf":
            return self._read_pdf(content_bytes)
        elif ext in ("docx", "doc"):
            return self._read_docx(content_bytes)
        else:
            try:
                return self._read_text(content_bytes)
            except Exception:
                return {"error": f"Unsupported file type: .{ext}"}

    # ------------------------------------------------------------------
    # File type parsers
    # ------------------------------------------------------------------

    def _read_text(self, content_bytes):
        text = content_bytes.decode("utf-8", errors="replace")
        if len(text) > 10000:
            return {
                "type": "text",
                "content": text[:10000],
                "truncated": True,
                "total_chars": len(text),
            }
        return {"type": "text", "content": text, "truncated": False}

    def _read_json(self, content_bytes):
        import json

        text = content_bytes.decode("utf-8", errors="replace")
        try:
            data = json.loads(text)
            if isinstance(data, list) and data and isinstance(data[0], dict):
                return {
                    "type": "json_table",
                    "headers": list(data[0].keys()),
                    "rows": data[:100],
                    "total_rows": len(data),
                }
            return {"type": "json", "content": data}
        except json.JSONDecodeError:
            return {"type": "text", "content": text[:10000]}

    def _read_csv(self, content_bytes, max_rows=100):
        import csv
        import io

        text = content_bytes.decode("utf-8", errors="replace")
        reader = csv.reader(io.StringIO(text))
        rows = []
        headers = None
        for i, row in enumerate(reader):
            if i == 0:
                headers = row
            elif i <= max_rows:
                rows.append(row)
        return {
            "type": "csv",
            "headers": headers,
            "rows": rows,
            "total_rows": len(rows),
            "sample": _format_as_text_table(headers, rows[:20]),
        }

    def _read_excel(self, content_bytes, max_rows=100):
        try:
            import io

            import openpyxl

            wb = openpyxl.load_workbook(
                io.BytesIO(content_bytes), read_only=True, data_only=True
            )
            sheets = {}
            for sheet_name in wb.sheetnames[:5]:  # Max 5 sheets
                ws = wb[sheet_name]
                rows = []
                headers = None
                for i, row in enumerate(ws.iter_rows(values_only=True)):
                    if i == 0:
                        headers = [
                            str(c) if c else f"Col{j}" for j, c in enumerate(row)
                        ]
                    elif i <= max_rows:
                        rows.append(
                            [str(c) if c is not None else "" for c in row]
                        )
                sheets[sheet_name] = {
                    "headers": headers,
                    "rows": rows,
                    "total_rows": ws.max_row - 1 if ws.max_row else 0,
                    "sample": _format_as_text_table(headers, rows[:20]),
                }
            wb.close()
            return {"type": "excel", "sheets": sheets}
        except ImportError:
            return {
                "error": "openpyxl not installed. Ask admin to run: pip install openpyxl"
            }

    def _read_pdf(self, content_bytes):
        # Try PyPDF2 first
        try:
            import io

            import PyPDF2

            reader = PyPDF2.PdfReader(io.BytesIO(content_bytes))
            text = ""
            for i, page in enumerate(reader.pages[:20]):  # Max 20 pages
                page_text = page.extract_text() or ""
                text += f"\n--- Page {i+1} ---\n{page_text}"
            text = text.strip()
            if not text:
                return {
                    "type": "pdf",
                    "content": "(PDF has no extractable text — may be scanned/image-based)",
                    "pages": len(reader.pages),
                }
            if len(text) > 10000:
                return {
                    "type": "pdf",
                    "content": text[:10000],
                    "truncated": True,
                    "pages": len(reader.pages),
                }
            return {"type": "pdf", "content": text, "pages": len(reader.pages)}
        except ImportError:
            pass

        # Try pdfplumber as fallback
        try:
            import io

            import pdfplumber

            text = ""
            with pdfplumber.open(io.BytesIO(content_bytes)) as pdf:
                for i, page in enumerate(pdf.pages[:20]):
                    page_text = page.extract_text() or ""
                    text += f"\n--- Page {i+1} ---\n{page_text}"
            text = text.strip()
            if len(text) > 10000:
                return {"type": "pdf", "content": text[:10000], "truncated": True}
            return {"type": "pdf", "content": text}
        except ImportError:
            return {
                "error": "No PDF reader installed. Ask admin to run: pip install PyPDF2"
            }

    def _read_docx(self, content_bytes):
        """Extract text from Word documents (.docx)."""
        try:
            import docx
            import io

            doc = docx.Document(io.BytesIO(content_bytes))
            text = ""

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"

            # Extract tables
            for table in doc.tables:
                text += "\n"
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    text += " | ".join(cells) + "\n"
                text += "\n"

            text = text.strip()
            if not text:
                return {
                    "type": "docx",
                    "content": "(Document appears empty or contains only images)",
                }
            if len(text) > 10000:
                return {
                    "type": "docx",
                    "content": text[:10000],
                    "truncated": True,
                    "total_chars": len(text),
                }
            return {"type": "docx", "content": text, "truncated": False}
        except ImportError:
            return {
                "error": "python-docx not installed. Ask admin to run: pip install python-docx"
            }

    # ------------------------------------------------------------------
    # File reading helper (same pattern as analyze_image)
    # ------------------------------------------------------------------

    def _get_file_content(self, file_url):
        """Read file content as bytes from Frappe."""
        import os

        try:
            # Try Frappe File doctype first
            files = frappe.get_all(
                "File",
                filters={"file_url": file_url},
                fields=["name"],
                limit_page_length=1,
            )
            if files:
                doc = frappe.get_doc("File", files[0]["name"])
                content = doc.get_content()
                if content:
                    if isinstance(content, str):
                        content = content.encode("utf-8")
                    return content

            # Fallback: direct path
            if file_url.startswith("/"):
                filename = (
                    file_url.split("/files/")[-1]
                    if "/files/" in file_url
                    else ""
                )
                if filename:
                    for prefix in ("private", "public"):
                        path = frappe.utils.get_site_path(prefix, "files", filename)
                        if os.path.exists(path):
                            with open(path, "rb") as f:
                                return f.read()
        except Exception as e:
            frappe.log_error(
                title="File read failed", message=f"{file_url}: {e}"
            )
        return None


# ------------------------------------------------------------------
# Module-level helper
# ------------------------------------------------------------------

def _format_as_text_table(headers, rows):
    """Format rows as a readable markdown-style text table for the AI."""
    if not headers or not rows:
        return ""
    lines = [" | ".join(str(h) for h in headers)]
    lines.append(" | ".join(["---"] * len(headers)))
    for row in rows:
        padded = list(row) + [""] * (len(headers) - len(row))
        lines.append(" | ".join(str(c) for c in padded[: len(headers)]))
    return "\n".join(lines)
